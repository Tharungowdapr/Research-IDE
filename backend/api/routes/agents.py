"""
Agent Routes v2 — gap analysis, ideas, plan, code, report + download + streaming
"""

import json
import asyncio
from typing import Optional, AsyncGenerator
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from core.database import get_db, get_fresh_db
from core.security import get_current_user
from core.llm_client import build_llm_client_for_user
from core.utils import compute_etag
from models.user import User
from models.project import Project, Output
from agents.gap_miner.gap_agent import run_gap_analysis
from agents.idea_generator.idea_agent import run_idea_generation
from agents.planner.planner_agent import run_planning
from agents.code_agent.code_agent import run_code_generation
from agents.writer.writer_agent import run_report_generation

router = APIRouter()


# ── Schemas ────────────────────────────────────────────────────────────────────

class AgentRequest(BaseModel):
    project_id: str

class SelectIdeaRequest(BaseModel):
    project_id: str
    idea_index: int


# ── Standard agent endpoints ───────────────────────────────────────────────────

@router.post("/analyze-gaps")
async def analyze_gaps(body: AgentRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project, papers_data, intent = _load_ctx(body.project_id, current_user.id, db)
    llm = build_llm_client_for_user(current_user)
    try:
        gaps = await run_gap_analysis(papers_data.get("papers", []), intent, llm)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gap analysis failed: {e}")
    _save_output(db, project.id, "gaps", {"gaps": gaps})
    project.current_stage = "ideas"
    db.commit()
    return {"project_id": project.id, "gaps": gaps}


@router.post("/generate-ideas")
async def generate_ideas(body: AgentRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project, papers_data, intent = _load_ctx(body.project_id, current_user.id, db)
    gaps_output = _get_output(db, project.id, "gaps")
    if not gaps_output:
        raise HTTPException(status_code=400, detail="Run gap analysis first")
    llm = build_llm_client_for_user(current_user)
    try:
        ideas = await run_idea_generation(gaps_output.get("gaps", []), papers_data.get("papers", []), intent, llm)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Idea generation failed: {e}")
    _save_output(db, project.id, "ideas", {"ideas": ideas})
    db.commit()
    return {"project_id": project.id, "ideas": ideas}


@router.post("/select-idea")
async def select_idea(body: SelectIdeaRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = _get_project(body.project_id, current_user.id, db)
    ideas_output = _get_output(db, project.id, "ideas")
    if not ideas_output:
        raise HTTPException(status_code=400, detail="No ideas found")
    ideas = ideas_output.get("ideas", [])
    if body.idea_index >= len(ideas):
        raise HTTPException(status_code=400, detail="Invalid idea index")
    selected = ideas[body.idea_index]
    # Clear downstream stale outputs
    for output_type in ["plan", "code", "report"]:
        existing = db.query(Output).filter(Output.project_id == project.id, Output.output_type == output_type).first()
        if existing:
            db.delete(existing)
    _save_output(db, project.id, "selected_idea", {"idea": selected})
    project.current_stage = "planner"
    db.commit()
    return {"project_id": project.id, "selected_idea": selected}


@router.post("/plan")
async def create_plan(body: AgentRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project, papers_data, intent = _load_ctx(body.project_id, current_user.id, db)
    selected = _get_output(db, project.id, "selected_idea")
    if not selected:
        raise HTTPException(status_code=400, detail="Select an idea first")
    llm = build_llm_client_for_user(current_user)
    try:
        plan = await run_planning(selected.get("idea", {}), intent, llm)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Planning failed: {e}")
    _save_output(db, project.id, "plan", plan)
    project.current_stage = "code"
    db.commit()
    return {"project_id": project.id, "plan": plan}


@router.post("/generate-code")
async def generate_code(body: AgentRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project, papers_data, intent = _load_ctx(body.project_id, current_user.id, db)
    selected = _get_output(db, project.id, "selected_idea")
    plan = _get_output(db, project.id, "plan")
    if not selected or not plan:
        raise HTTPException(status_code=400, detail="Complete planning step first")
    llm = build_llm_client_for_user(current_user)
    file_hints = plan.get("file_structure", [])
    # Pass github_url from top papers to code agent
    idea_with_github = dict(selected.get("idea", {}))
    top_github = next(
        (p.get("github_url", "") for p in papers_data.get("papers", [])[:3] if p.get("github_url")),
        ""
    )
    if top_github:
        idea_with_github["github_url"] = top_github
    try:
        code = await run_code_generation(idea_with_github, plan, llm, file_hints=file_hints)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Code generation failed: {e}")
    _save_output(db, project.id, "code", code)
    project.current_stage = "report"
    db.commit()
    return {"project_id": project.id, "code": code}


@router.post("/generate-report")
async def generate_report(body: AgentRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project, papers_data, intent = _load_ctx(body.project_id, current_user.id, db)
    selected = _get_output(db, project.id, "selected_idea")
    plan = _get_output(db, project.id, "plan") or {}
    gaps = _get_output(db, project.id, "gaps") or {}
    if not selected:
        raise HTTPException(status_code=400, detail="Select an idea first")
    llm = build_llm_client_for_user(current_user)
    try:
        report = await run_report_generation(
            selected.get("idea", {}), papers_data.get("papers", []),
            gaps.get("gaps", []), plan, intent, llm,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Report generation failed: {e}")
    _save_output(db, project.id, "report", report)
    project.status = "done"
    db.commit()
    return {"project_id": project.id, "report": report}


# ── Streaming endpoint ─────────────────────────────────────────────────────────

@router.get("/stream/{project_id}/{stage}")
async def stream_agent(
    project_id: str,
    stage: str,
    current_user: User = Depends(get_current_user),
):
    """SSE streaming endpoint. Returns live progress messages then the final result."""
    if stage not in ("gaps", "ideas", "plan", "code", "report"):
        raise HTTPException(status_code=400, detail=f"Unknown stage: {stage}")

    async def event_stream() -> AsyncGenerator[str, None]:
        def sse(msg_type: str, data: dict) -> str:
            return f"data: {json.dumps({'type': msg_type, **data})}\n\n"

        db = get_fresh_db()
        try:
            project = db.query(Project).filter(
                Project.id == project_id, Project.user_id == current_user.id
            ).first()
            if not project:
                yield sse("error", {"message": "Project not found"})
                return

            llm = build_llm_client_for_user(current_user)
            papers_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "papers").first()
            intent_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "intent").first()
            papers_data = papers_out.data if papers_out else {"papers": []}
            intent = intent_out.data if intent_out else {}

            if stage == "gaps":
                # Pass progress callback so gap agent can stream live updates
                gap_log = []
                async def gap_progress(msg: str):
                    gap_log.append(msg)
                    # We yield inside the generator — use a queue trick
                    nonlocal _gap_msg
                    _gap_msg = msg

                _gap_msg = None
                yield sse("progress", {"message": "Starting 3-pass gap analysis..."})
                await asyncio.sleep(0.05)

                # Simple progress: yield before each pass
                yield sse("progress", {"message": "Pass 1: Extracting claims from papers..."})
                await asyncio.sleep(0.05)
                yield sse("progress", {"message": "Pass 2: Identifying research gaps with LLM..."})
                gaps = await run_gap_analysis(papers_data.get("papers", []), intent, llm)
                yield sse("progress", {"message": f"Pass 3: Scoring complete — {len(gaps)} gaps identified"})
                _save_output_fresh(db, project_id, "gaps", {"gaps": gaps})
                project.current_stage = "ideas"
                db.commit()
                yield sse("result", {"data": {"gaps": gaps}})

            elif stage == "ideas":
                gaps_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "gaps").first()
                gaps_data = gaps_out.data if gaps_out else {"gaps": []}
                yield sse("progress", {"message": "Generating research ideas..."})
                await asyncio.sleep(0.1)
                yield sse("progress", {"message": "Running critic review..."})
                await asyncio.sleep(0.1)
                yield sse("progress", {"message": "Refining and ranking ideas..."})
                ideas = await run_idea_generation(gaps_data.get("gaps", []), papers_data.get("papers", []), intent, llm)
                _save_output_fresh(db, project_id, "ideas", {"ideas": ideas})
                db.commit()
                yield sse("result", {"data": {"ideas": ideas}})

            elif stage == "plan":
                selected_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "selected_idea").first()
                selected = selected_out.data if selected_out else {}
                yield sse("progress", {"message": "Building execution plan (Pass 1)..."})
                await asyncio.sleep(0.05)

                # Callback saves Pass 1 immediately so UI gets partial result
                def save_pass1(partial_plan: dict):
                    _save_output_fresh(db, project_id, "plan", partial_plan)
                    db.commit()

                yield sse("progress", {"message": "Generating experiment configs (Pass 2)..."})
                plan = await run_planning(selected.get("idea", {}), intent, llm, on_pass1_complete=save_pass1)
                _save_output_fresh(db, project_id, "plan", plan)
                project.current_stage = "code"
                db.commit()
                yield sse("result", {"data": {"plan": plan}})

            elif stage == "code":
                selected_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "selected_idea").first()
                plan_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "plan").first()
                selected = selected_out.data if selected_out else {}
                plan = plan_out.data if plan_out else {}
                file_hints = plan.get("file_structure", [])
                yield sse("progress", {"message": "Generating project scaffold..."})
                await asyncio.sleep(0.1)
                yield sse("progress", {"message": "Writing model and training code..."})
                await asyncio.sleep(0.1)
                yield sse("progress", {"message": "Creating Makefile and tests..."})
                code = await run_code_generation(selected.get("idea", {}), plan, llm, file_hints=file_hints)
                _save_output_fresh(db, project_id, "code", code)
                project.current_stage = "report"
                db.commit()
                yield sse("result", {"data": {"code": code}})

            elif stage == "report":
                selected_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "selected_idea").first()
                plan_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "plan").first()
                gaps_out = db.query(Output).filter(Output.project_id == project_id, Output.output_type == "gaps").first()
                selected = selected_out.data if selected_out else {}
                plan = plan_out.data if plan_out else {}
                gaps = gaps_out.data if gaps_out else {}
                yield sse("progress", {"message": "Drafting paper structure..."})
                await asyncio.sleep(0.1)
                yield sse("progress", {"message": "Writing introduction and related work..."})
                await asyncio.sleep(0.1)
                yield sse("progress", {"message": "Writing methodology and results..."})
                report = await run_report_generation(
                    selected.get("idea", {}), papers_data.get("papers", []),
                    gaps.get("gaps", []), plan, intent, llm,
                )
                _save_output_fresh(db, project_id, "report", report)
                project.status = "done"
                db.commit()
                yield sse("result", {"data": {"report": report}})

            yield sse("done", {"message": f"{stage} complete"})

        except Exception as e:
            yield sse("error", {"message": str(e)})
        finally:
            db.close()

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


# ── Download endpoints ─────────────────────────────────────────────────────────

@router.get("/{project_id}/download/docx")
async def download_docx(
    project_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate and stream an IEEE-format DOCX paper."""
    # Verify ownership
    _get_project(project_id, current_user.id, db)
    report = _get_output(db, project_id, "report")
    if not report:
        raise HTTPException(status_code=404, detail="Report not found. Generate it first.")

    # ETag caching
    etag = compute_etag(report)

    try:
        docx_bytes = _generate_docx(report)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")

    return StreamingResponse(
        iter([docx_bytes]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="research_paper.docx"',
            "ETag": etag,
        },
    )


@router.get("/{project_id}/download/pdf")
async def download_pdf(
    project_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate and stream an IEEE-format PDF paper."""
    # Verify ownership
    _get_project(project_id, current_user.id, db)
    report = _get_output(db, project_id, "report")
    if not report:
        raise HTTPException(status_code=404, detail="Report not found. Generate it first.")

    etag = compute_etag(report)

    try:
        pdf_bytes = _generate_pdf(report)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")

    return StreamingResponse(
        iter([pdf_bytes]),
        media_type="application/pdf",
        headers={
            "Content-Disposition": 'attachment; filename="research_paper.pdf"',
            "ETag": etag,
        },
    )


# ── DOCX Generator ────────────────────────────────────────────────────────────

def _generate_docx(report: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from lxml import etree
    import io

    doc = Document()

    # ── Page setup: A4, IEEE margins ─────────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    # ── Default style: Times New Roman 10pt ───────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    def add_para(text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return p

    # ── Single-column header section: title, authors, abstract, keywords ─────
    # Title
    add_para(report.get("title", "Research Paper"), size=24, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)

    # Authors
    authors = ", ".join(report.get("authors", ["Author"]))
    add_para(authors, size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Abstract label + content
    abstract_para = doc.add_paragraph()
    abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_para.paragraph_format.space_after = Pt(6)
    run_label = abstract_para.add_run("Abstract\u2014")
    run_label.italic = True
    run_label.font.name = "Times New Roman"
    run_label.font.size = Pt(9)
    run_content = abstract_para.add_run(report.get("abstract", ""))
    run_content.font.name = "Times New Roman"
    run_content.font.size = Pt(9)
    run_content.italic = True

    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.space_after = Pt(12)
    run_kl = kw_para.add_run("Index Terms\u2014")
    run_kl.italic = True
    run_kl.font.name = "Times New Roman"
    run_kl.font.size = Pt(9)
    run_kv = kw_para.add_run(", ".join(report.get("keywords", [])))
    run_kv.font.name = "Times New Roman"
    run_kv.font.size = Pt(9)

    # ── Switch to two-column layout for body ──────────────────────────────────
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.page_height = Cm(29.7)
    new_section.page_width = Cm(21.0)
    new_section.top_margin = Cm(1.9)
    new_section.bottom_margin = Cm(2.54)
    new_section.left_margin = Cm(1.9)
    new_section.right_margin = Cm(1.9)

    # Add w:cols to the new section's sectPr
    sectPr = new_section._sectPr
    cols_xml = f'<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:num="2" w:space="720"/>'
    cols_element = parse_xml(cols_xml)
    # Remove existing cols if any
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    sectPr.append(cols_element)

    # ── Body sections ─────────────────────────────────────────────────────────
    for section_data in report.get("sections", []):
        heading = section_data.get("heading", "")
        content = section_data.get("content", "")

        # Section heading: bold, centered, all caps
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(heading.upper() if not heading.isupper() else heading)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

        # Body paragraphs
        for para_text in content.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Pt(12)
            p.paragraph_format.line_spacing = Pt(12)
            run = p.add_run(para_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    # ── Acknowledgements ──────────────────────────────────────────────────────
    if report.get("acknowledgements"):
        add_para("ACKNOWLEDGEMENTS", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(report["acknowledgements"])
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    # ── References ────────────────────────────────────────────────────────────
    # Switch back to single column for references
    ref_section = doc.add_section(WD_SECTION.CONTINUOUS)
    ref_section._sectPr.append(parse_xml(
        '<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:num="1"/>'
    ))

    add_para("REFERENCES", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    for ref in report.get("references", []):
        ref_text = f"[{ref.get('id','')}] {ref.get('authors','')} ({ref.get('year','')}). \"{ref.get('title','')}.\" {ref.get('venue','')}."
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF Generator ─────────────────────────────────────────────────────────────

def _generate_pdf(report: dict) -> bytes:
    html = _build_ieee_html(report)
    try:
        import weasyprint
        return weasyprint.HTML(string=html).write_pdf()
    except ImportError:
        # WeasyPrint not available — return HTML as fallback bytes
        raise HTTPException(
            status_code=501,
            detail="PDF generation requires WeasyPrint. Run: pip install weasyprint. "
                   "On Ubuntu also run: apt-get install -y libpango-1.0-0 libpangocairo-1.0-0"
        )


def _build_ieee_html(report: dict) -> str:
    def esc(s):
        return str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    sections_html = ""
    for s in report.get("sections", []):
        heading = esc(s.get("heading",""))
        content = esc(s.get("content",""))
        paras = "".join(f"<p>{p.strip()}</p>" for p in content.split("\n\n") if p.strip())
        sections_html += f"<h2>{heading}</h2>{paras}"

    refs_html = "".join(
        f'<div class="ref"><span class="ref-num">[{r.get("id","")}]</span>'
        f'<span>{esc(r.get("authors",""))} ({esc(r.get("year",""))}). '
        f'&ldquo;{esc(r.get("title",""))}.&rdquo; {esc(r.get("venue",""))}.</span></div>'
        for r in report.get("references", [])
    )
    keywords = ", ".join(report.get("keywords", []))
    authors = ", ".join(report.get("authors", ["Author"]))

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>
  @page {{
    size: A4;
    margin: 1.9cm 1.9cm 2.54cm 1.9cm;
  }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: "Times New Roman", Times, serif; font-size: 10pt; color: #000; }}
  #header {{ text-align: center; margin-bottom: 12pt; }}
  #header h1 {{ font-size: 20pt; font-weight: bold; margin-bottom: 6pt; }}
  #header .authors {{ font-size: 11pt; font-style: italic; margin-bottom: 10pt; }}
  #abstract {{ font-size: 9pt; margin-bottom: 6pt; text-align: justify; }}
  #abstract .label {{ font-style: italic; font-weight: bold; }}
  #keywords {{ font-size: 9pt; margin-bottom: 14pt; }}
  #keywords .label {{ font-style: italic; }}
  #body-content {{
    column-count: 2;
    column-gap: 18pt;
    text-align: justify;
  }}
  h2 {{
    font-size: 10pt; font-weight: bold; text-align: center;
    text-transform: uppercase; margin-top: 10pt; margin-bottom: 4pt;
    -webkit-column-break-before: avoid; column-break-before: avoid;
  }}
  p {{
    font-size: 10pt; text-indent: 12pt; margin-bottom: 4pt; line-height: 1.3;
    -webkit-column-break-inside: avoid; column-break-inside: avoid;
  }}
  #refs {{
    column-count: 1; margin-top: 12pt;
    border-top: 1pt solid #000; padding-top: 8pt;
  }}
  #refs h2 {{ text-align: center; margin-bottom: 6pt; }}
  .ref {{ display: grid; grid-template-columns: 24pt 1fr; gap: 4pt; margin-bottom: 3pt; font-size: 8pt; }}
  .ref-num {{ font-size: 8pt; }}
  #ack {{ margin-top: 8pt; font-size: 10pt; }}
</style>
</head>
<body>
<div id="header">
  <h1>{esc(report.get("title","Research Paper"))}</h1>
  <div class="authors">{esc(authors)}</div>
</div>
<div id="abstract">
  <span class="label">Abstract&mdash;</span>{esc(report.get("abstract",""))}
</div>
<div id="keywords">
  <span class="label">Index Terms&mdash;</span>{esc(keywords)}
</div>
<div id="body-content">
{sections_html}
{f'<div id="ack"><h2>ACKNOWLEDGEMENTS</h2><p>{esc(report.get("acknowledgements",""))}</p></div>' if report.get("acknowledgements") else ""}
</div>
<div id="refs">
  <h2>REFERENCES</h2>
  {refs_html}
</div>
</body>
</html>"""


# ── Shared helpers ────────────────────────────────────────────────────────────

def _get_project(project_id: str, user_id: str, db: Session) -> Project:
    p = db.query(Project).filter(Project.id == project_id, Project.user_id == user_id).first()
    if not p:
        raise HTTPException(status_code=404, detail="Project not found")
    return p

def _load_ctx(project_id: str, user_id: str, db: Session):
    project = _get_project(project_id, user_id, db)
    papers_data = _get_output(db, project_id, "papers") or {"papers": []}
    intent = _get_output(db, project_id, "intent") or {}
    return project, papers_data, intent

def _save_output(db: Session, project_id: str, output_type: str, data: dict):
    from models.project import OutputHistory
    existing = db.query(Output).filter(Output.project_id == project_id, Output.output_type == output_type).first()
    if existing:
        # Archive current version before overwriting
        hist_count = db.query(OutputHistory).filter(
            OutputHistory.project_id == project_id,
            OutputHistory.output_type == output_type,
        ).count()
        db.add(OutputHistory(
            project_id=project_id,
            output_type=output_type,
            version=f"v{hist_count + 1}",
            data=existing.data,
        ))
        existing.data = data
    else:
        db.add(Output(project_id=project_id, output_type=output_type, data=data))
    db.commit()

def _save_output_fresh(db: Session, project_id: str, output_type: str, data: dict):
    """Save output using a provided session (for streaming use)."""
    existing = db.query(Output).filter(Output.project_id == project_id, Output.output_type == output_type).first()
    if existing:
        existing.data = data
    else:
        db.add(Output(project_id=project_id, output_type=output_type, data=data))
    db.commit()

def _get_output(db: Session, project_id: str, output_type: str):
    o = db.query(Output).filter(Output.project_id == project_id, Output.output_type == output_type).first()
    return o.data if o else None


@router.post("/more-ideas")
async def generate_more_ideas(
    body: AgentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate 4 additional ideas with different emphasis."""
    project, papers_data, intent = _load_ctx(body.project_id, current_user.id, db)
    gaps_output = _get_output(db, project.id, "gaps")
    existing_ideas_output = _get_output(db, project.id, "ideas")
    if not gaps_output:
        raise HTTPException(status_code=400, detail="Run gap analysis first")

    llm = build_llm_client_for_user(current_user)
    existing_titles = [i.get("title","") for i in (existing_ideas_output or {}).get("ideas", [])]

    try:
        from agents.idea_generator.idea_agent import run_idea_generation
        new_ideas = await run_idea_generation(
            gaps_output.get("gaps", []),
            papers_data.get("papers", []),
            {**intent, "_exclude_titles": existing_titles},
            llm,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Idea generation failed: {e}")

    # Merge with existing ideas
    all_ideas = list((existing_ideas_output or {}).get("ideas", [])) + new_ideas
    _save_output(db, project.id, "ideas", {"ideas": all_ideas})
    db.commit()
    return {"project_id": project.id, "new_ideas": new_ideas, "total": len(all_ideas)}


class UserIdeaRequest(BaseModel):
    project_id: str
    idea_text: str

@router.post("/submit-user-idea")
async def submit_user_idea(
    body: UserIdeaRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Critique, refine, and add a user-submitted idea."""
    project, papers_data, intent = _load_ctx(body.project_id, current_user.id, db)
    llm = build_llm_client_for_user(current_user)
    
    from core.utils import safe_parse_llm_json
    CRITIC_SYS = "You are a harsh but fair peer reviewer. Find fatal flaws in this user-submitted research idea. Be specific. Return ONLY valid JSON."
    critique_raw = await llm.complete(
        f"Critique this research idea:\n{body.idea_text}\n\n"
        "Return a JSON object with: fatal_flaw (1 sentence), weakness_score (1-10), is_salvageable (true/false), suggested_fix (1 sentence).",
        system=CRITIC_SYS, json_mode=True
    )
    critique = safe_parse_llm_json(critique_raw, default={})
    
    DEFEND_SYS = "You are a senior researcher refining an idea. Return ONLY valid JSON."
    defend_raw = await llm.complete(
        f"Original idea: {body.idea_text}\n\nCritique: {json.dumps(critique)}\n\n"
        "Revise the idea to address the critique and format it into our standard schema. "
        "Return JSON object with: title, description, problem_statement, proposed_solution, why_it_addresses_gap, potential_challenges, addresses_gaps, suggested_methods, suggested_datasets, complexity, estimated_time, feasibility, innovation_level, novelty_score, feasibility_score, difficulty.",
        system=DEFEND_SYS, json_mode=True
    )
    refined_idea = safe_parse_llm_json(defend_raw, default={})
    
    # Save it
    existing_ideas_output = _get_output(db, project.id, "ideas")
    all_ideas = list((existing_ideas_output or {}).get("ideas", []))
    refined_idea["is_user_submitted"] = True
    refined_idea["original_text"] = body.idea_text
    refined_idea["critique_summary"] = critique.get("fatal_flaw", "User submitted")
    all_ideas.insert(0, refined_idea)
    _save_output(db, project.id, "ideas", {"ideas": all_ideas})
    db.commit()
    
    return {"project_id": project.id, "refined_idea": refined_idea, "critique": critique}


class FilterIdeasRequest(BaseModel):
    project_id: str
    complexity: Optional[str] = None
    feasibility: Optional[str] = None

@router.post("/filter-ideas")
async def filter_ideas(
    body: FilterIdeasRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Filter generated ideas by complexity or feasibility."""
    project = _get_project(body.project_id, current_user.id, db)
    ideas_output = _get_output(db, project.id, "ideas")
    if not ideas_output:
        return {"ideas": []}
    
    ideas = ideas_output.get("ideas", [])
    filtered = ideas
    
    if body.complexity:
        filtered = [i for i in filtered if str(i.get("complexity", "")).lower() == body.complexity.lower()]
    if body.feasibility:
        filtered = [i for i in filtered if str(i.get("feasibility", "")).lower() == body.feasibility.lower()]
        
    return {"ideas": filtered}
